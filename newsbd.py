from bs4 import BeautifulSoup
import requests
from docx import Document
from docx.shared import Inches
from bitlyshortener import Shortener
import dailystarr
import datetime


topic = input("Please choose a topic \n 1. Economy \n 2. International \n 3. Technology \n 4. Job \n 4. Business \n")
if topic == '1':
    topic = 'economy'
elif topic == '2':
    topic = 'international'
elif topic == '3':
    topic = 'technology'
elif topic == '4':
    topic = 'chakri-bakri'
elif topic == '4':
    topic = 'business'
    execfile('dailystarr.py')



prothomturl = 'https://www.prothomalo.com/' + topic + '/'
paurl = 'https://www.prothomalo.com'
res = requests.get(prothomturl)
document = Document()
soup = BeautifulSoup(res.text, 'html.parser')

#print(paurl + soup.select('.content_type_article')[0].a.get('href'))
#prothomalo
articles = soup.select('.content_type_article')
def palo():
    title = []
    link = []
    for art in articles:
        headline = art.h2.text
        perm = art.a.get('href')
        perm = paurl + perm
        title.append(headline)
        link.append(perm)
    return title, link

title, link = palo()

tokens_pool = ['cc1e818a5979fe47c9aa975f5c8228f5e68692b4']
shortener = Shortener(tokens=tokens_pool, max_cache_size=10000)
short_link = shortener.shorten_urls(link)

s_link = []
for s in short_link:
    bit = s.replace('j.mp', 'bit.ly')
    s_link.append(bit)


document.add_heading('{}'.format(topic).upper(), 0)


for n,t in enumerate(title):
    document.add_heading(t, level=1)
    document.add_paragraph(s_link[n])
    
date = datetime.date.today()
document.save('{}-{}.docx'.format(topic, date))

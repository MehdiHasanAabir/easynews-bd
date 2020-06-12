from bs4 import BeautifulSoup
import requests
import re
from docx import Document
from docx.shared import Inches
import datetime

link = 'https://www.thedailystar.net/business'
ds = 'https://www.thedailystar.net'
thelink = requests.get(link)
soup = BeautifulSoup(thelink.text, 'html.parser')
document = Document()


art1 = soup.select('h2')
art3 = soup.select('h3')
art5 = soup.select('h5')



news = {}
news.update({art3[0].a.text:art3[0].a.get('href')})
for a3 in art3:
    title = a3.a.text
    link = ds + a3.a.get('href')
    news.update({title : link})
    
for a5 in art5:
    try:
        title = a5.a.text
        link = ds + a5.a.get('href')
    except:
        continue
    news.update({title : link})
   


#print(news)
art_num = str(len(news))
document.add_heading('Daily Star Business', 0)
document.add_paragraph(art_num + ' articles')


for n, l in news.items():
    document.add_heading(n, level=1)
    document.add_paragraph(l)
    
date = datetime.date.today()
document.save('Business-{}.docx'.format(date))
